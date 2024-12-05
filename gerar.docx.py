from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn



doc = Document()

# Função para estilizar tabelas
def estilizar_tabela(tabela):
    for linha in tabela.rows:
        for celula in linha.cells:
            celula.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            celula.paragraphs[0].style.font.size = Pt(11)
            celula.paragraphs[0].style.font.name = "Calibri"

# Adicionar título centralizado
titulo = doc.add_paragraph("MATRIZ DE PLANEJAMENTO E DESIGN EDUCACIONAL")
titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
titulo.style.font.size = Pt(14)
titulo.style.font.name = "Calibri"

# 1. DADOS GERAIS
doc.add_paragraph("\n1. DADOS GERAIS", style="Heading 2")
dados_gerais = doc.add_table(rows=6, cols=2)
dados_gerais.style = 'Table Grid'
conteudo_dados_gerais = [
    ["Curso", "CURSO TÉCNICO DE EVENTOS"],
    ["Disciplina", "ASPECTOS SOCIOCULTURAIS EM EVENTOS"],
    ["Semestre", "2024.2"],
    ["Período letivo de planejamento", "2024.1"],
    ["Formato de oferta da disciplina", "MODULAR"],
    ["Professor(a)", "CAROLINA CASTELO BRANCO"]
]
for i, (col1, col2) in enumerate(conteudo_dados_gerais):
    row = dados_gerais.rows[i]
    row.cells[0].text = col1
    row.cells[1].text = col2

doc.add_paragraph("\n2. DISTRIBUIÇÃO DE CARGA HORÁRIA E UNIDADES", style="Heading 2")

doc.add_paragraph("\n2.1 Distribuição da carga horária", style="Heading 3")
tabela_carga = doc.add_table(rows=2, cols=3)
tabela_carga.style = 'Table Grid'
tabela_carga.cell(0, 0).text = "Carga horária total da disciplina: 40h"
tabela_carga.cell(0, 1).text = "Carga horária a distância: 36h"
tabela_carga.cell(0, 2).text = "Carga horária presencial: 4h"
tabela_carga.cell(1, 0).text = "CH síncrona (meet): 8h"
tabela_carga.cell(1, 1).text = "CH assíncrona: 28h"
tabela_carga.cell(1, 2).text = "Total de unidades da disciplina: 4"


doc.add_paragraph("\n2.2 Distribuição das unidades", style="Heading 3")
tabela_unidades = doc.add_table(rows=5, cols=4) 
tabela_unidades.style = 'Table Grid'
unidades = [
    ["Unidade", "CH do semestre", "Título das etapas", "Período"],
    ["1", "10h", "Histórico, Conceitos e Aspectos Socioculturais dos Eventos", "19/09/2024 a 25/09/2024"],
    ["2", "10h", "Cultura Local e os Eventos", "26/09/2024 a 02/10/2024"],
    ["3", "10h", "Responsabilidade Social e Ética nos Eventos", "03/10/2024 a 09/10/2024"],
    ["4", "10h", "Sustentabilidade nos Eventos", "10/10/2024 a 16/10/2024"]
]
for i, row_content in enumerate(unidades):
    row = tabela_unidades.rows[i]
    for j, cell_content in enumerate(row_content):
        row.cells[j].text = cell_content


doc.add_paragraph("\n2.3 Agenda de atividades síncronas", style="Heading 3")
tabela_agenda = doc.add_table(rows=2, cols=3)
tabela_agenda.style = 'Table Grid'
agenda = [
    ["Data", "CH dos encontros presenciais", "Objetivos do meet/encontro presencial"],
    ["24/09/2024", "2h/a", "Encontro síncrono 1:\n- Apresentação da disciplina, do PUD e discussão das regras de convivência virtual e presencial;\n- Aula expositiva sobre o histórico, conceitos e Aspectos Socioculturais dos eventos."]
   
]
for i, row_content in enumerate(agenda):
    row = tabela_agenda.rows[i]
    for j, cell_content in enumerate(row_content):
        row.cells[j].text = cell_content
        
        
# 3. CALENDÁRIO/CRONOGRAMA DA DISCIPLINA
doc.add_paragraph(" CALENDÁRIO/CRONOGRAMA DA DISCIPLINA", style="Heading 2")
tabela_cronograma = doc.add_table(rows=1, cols=2)
tabela_cronograma.style = 'Table Grid'

# Preencher cronograma
cronograma = [
    ["UNIDADE 1 – 19/09/2024 a 25/09/2024", "Aula Síncrona 1 – 24/09/2024 – Google Meet (18h20min às 20h20min)"],
    ["UNIDADE 2 – 26/09/2024 a 02/10/2024", "Aula Síncrona 2 – 01/10/2024 – Google Meet (18h20min às 20h20min)"],
    ["UNIDADE 3 – 03/10/2024 a 09/10/2024", "Aula Síncrona 3 – 08/10/2024 – Google Meet (18h20min às 20h20min)\nENCONTRO PRESENCIAL 1 AVALIAÇÃO I – 09/10/2024 (20h às 22h00min)"],
    ["UNIDADE 4 – 10/10/2024 a 16/10/2024", "Aula Síncrona 4 – 15/10/2024 – Google Meet (18h20min às 20h20min)\nENCONTRO PRESENCIAL 2 AVALIAÇÃO II – 16/10/2024 (20h às 22h00min)"]
]
for linha in cronograma:
    row = tabela_cronograma.add_row()
    row.cells[0].text = linha[0]
    row.cells[1].text = linha[1]

# 4. DESCRIÇÃO DO MURAL
doc.add_paragraph("\n4. DESCRIÇÃO DO MURAL", style="Heading 2")
descricao = (
    "Olá alunos e alunas!\n\n"
    "Sejam bem-vindos(as) à disciplina ASPECTOS SOCIOCULTURAIS EM EVENTOS. Sou a Prof. Carolina "
    "Castelo Branco e estaremos juntos no decorrer deste primeiro semestre do curso Técnico em Eventos. "
    "Trabalharemos essa disciplina debatendo sobre o histórico e os conceitos do setor de eventos, com vistas a compreender a importância "
    "dos aspectos socioculturais auxiliando na organização e execução dos mesmos, levando-se em consideração a cultura local e a importância da "
    "Responsabilidade Social, a Ética e a Sustentabilidade como fatores fundamentais para o desenvolvimento desse setor. Entenderemos também aspectos que definem a "
    "importância dessa área dentro da atividade turística, principalmente debatendo aspectos sobre a sua sustentabilidade econômica, física e ambiental "
    "no desenvolvimento desse setor.\n\n"
    "Outro fator fundamental para o desenvolvimento do conteúdo programático desta disciplina refere-se ao entendimento da diversidade dos eventos, além da valorização da cultura e a pluralidade de manifestações existentes no Ceará. Devemos compreender que atualmente a Acessibilidade e Inclusão são fatores essenciais para se atuar com ética e responsabilidade social, como contributo para o desenvolvimento da sustentabilidade dos eventos socioculturais auxiliando no desenvolvimento turístico de uma região.\n\n"
    "As aulas síncronas serão nas terças, no horário 18:20 às 20:20. Teremos encontros síncronos nos dias 24/09, 01/10 e 08/10. Link para aula: https://meet.google.com/ffe-avcu-vkc\n\n"
    "Encontro Presencial dias 09/10 e 16/10, de 20h às 22h.\n\n"
    "● Recurso “Fórum de discussão” (Exemplo: Fórum tira-dúvidas)\n"
    "● Configurar a Biblioteca com os materiais da disciplina. (Material em anexo)\n\n"
    "1. Oliveira, M. Apostila Introdução à Eventos, EtecBrasil, 2010\n"
    "2. Coutinho, Helen. Apostila Organização de Eventos, EtecBrasil, 2010\n"
)

paragrafo = doc.add_paragraph(descricao)

# Adicionar uma borda ao parágrafo
def adicionar_borda_no_paragrafo(paragrafo):
    p = paragrafo._element
    pPr = p.find(qn("w:pPr"))  # Encontrar propriedades do parágrafo
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p.insert(0, pPr)
    
    # Criar borda
    pbdr = OxmlElement("w:pBdr")
    
    # Configurar as bordas (topo, esquerda, base, direita)
    for border_name in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")  # Linha simples
        border.set(qn("w:sz"), "12")      # Tamanho da borda
        border.set(qn("w:space"), "4")   # Espaçamento
        border.set(qn("w:color"), "000000")  # Cor preta
        pbdr.append(border)
    
    # Adicionar borda ao parágrafo
    pPr.append(pbdr)

# Aplicar a borda ao parágrafo criado
adicionar_borda_no_paragrafo(paragrafo)



doc.save("matriz_planejamento.docx")
