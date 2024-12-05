# import os.path

# from google.auth.transport.requests import Request
# from google.oauth2.credentials import Credentials
# from google_auth_oauthlib.flow import InstalledAppFlow
# from googleapiclient.discovery import build
# from googleapiclient.errors import HttpError

# # If modifying these scopes, delete the file token.json.
# SCOPES = ["https://www.googleapis.com/auth/spreadsheets"]



# def main():
  
#   creds = None
  
#   if os.path.exists("token.json"):
#     creds = Credentials.from_authorized_user_file("token.json", SCOPES)
#   # If there are no (valid) credentials available, let the user log in.
#   if not creds or not creds.valid:
#     if creds and creds.expired and creds.refresh_token:
#       creds.refresh(Request())
#     else:
#       flow = InstalledAppFlow.from_client_secrets_file(
#           "client_secret.json", SCOPES
#       )
#       creds = flow.run_local_server(port=0)
#     # Save the credentials for the next run
#     with open("token.json", "w") as token:
#       token.write(creds.to_json())

#   try:
#     service = build("sheets", "v4", credentials=creds)

#     # Call the Sheets API
#     sheet = service.spreadsheets()
#     result = (
#         sheet.values()
#         .get(spreadsheetId="1dkbgrFTgRQWNCj839u0Mlb2WJSFmMoFPVAO6Y-BNQHs", range="Form responses!A1:CC5")
#         .execute()
#     )
#     values = result.get("values", [])
#     print(values)
   
#   except HttpError as err:
#     print(err)


# if __name__ == "__main__":
#   main()



import os.path
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Configuração da API do Google Sheets
SCOPES = ["https://www.googleapis.com/auth/spreadsheets"]

def get_sheet_data(spreadsheet_id, range_name):
    """Obtém os dados da planilha Google Sheets."""
    creds = None
    if os.path.exists("token.json"):
        creds = Credentials.from_authorized_user_file("token.json", SCOPES)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file("client_secret.json", SCOPES)
            creds = flow.run_local_server(port=0)
        with open("token.json", "w") as token:
            token.write(creds.to_json())
    try:
        service = build("sheets", "v4", credentials=creds)
        sheet = service.spreadsheets()
        result = sheet.values().get(spreadsheetId='1dkbgrFTgRQWNCj839u0Mlb2WJSFmMoFPVAO6Y-BNQHs', range='Form responses!A1:CC5').execute()
        return result.get("values", [])
    except HttpError as err:
        print(f"Erro ao acessar a planilha: {err}")
        return None

# Função para adicionar borda em parágrafos
def adicionar_borda_no_paragrafo(paragrafo):
    p = paragrafo._element
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p.insert(0, pPr)
    pbdr = OxmlElement("w:pBdr")
    for border_name in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "12")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), "000000")
        pbdr.append(border)
    pPr.append(pbdr)

def criar_documento(dados):
    """Cria o documento Word preenchendo os dados dinâmicos."""
    doc = Document()
    
    # Título
    titulo = doc.add_paragraph("MATRIZ DE PLANEJAMENTO E DESIGN EDUCACIONAL")
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # DADOS GERAIS
    doc.add_paragraph("\n1. DADOS GERAIS", style="Heading 2")
    dados_gerais = doc.add_table(rows=len(dados), cols=2)
    dados_gerais.style = 'Table Grid'
    for i, linha in enumerate(dados):
      campo = linha[0] if len(linha) > 0 else "Campo vazio"
      valor = linha[1] if len(linha) > 1 else "Valor vazio"
      row = dados_gerais.rows[i]
      row.cells[0].text = campo
      row.cells[1].text = valor

    
    # Exemplo de texto dinâmico
    descricao = (
        f"Bem-vindo(a) ao curso {dados[0][1]} ministrado por {dados[-1][1]}. "
        "Esperamos que aproveite as aulas!"
    )
    paragrafo = doc.add_paragraph(descricao)
    adicionar_borda_no_paragrafo(paragrafo)
    
    doc.save("matriz_planejamento.docx")
    print("Documento gerado com sucesso!")

# Main
if __name__ == "__main__":
    # ID da planilha e intervalo
    SPREADSHEET_ID = "1dkbgrFTgRQWNCj839u0Mlb2WJSFmMoFPVAO6Y-BNQHs"
    RANGE_NAME = "Form responses!A1:CC5"
    
    # Obter dados da planilha
    dados_planilha = get_sheet_data(SPREADSHEET_ID, RANGE_NAME)
    if dados_planilha:
        # Remove o cabeçalho se necessário
        dados_planilha = dados_planilha[1:]
        criar_documento(dados_planilha)
